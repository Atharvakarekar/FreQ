# flask_to_streamlit_utils.py
import os
import json
import re
import docx
import pdfplumber
import pytesseract
# import SentenceTransformers
from PIL import Image
from collections import defaultdict
from sentence_transformers import SentenceTransformer, util
import nltk
from PyPDF2 import PdfReader, PdfWriter

import nltk
nltk.download('punkt', download_dir='C:/nltk_data')
from nltk.tokenize.punkt import PunktSentenceTokenizer
tokenizer = PunktSentenceTokenizer()

def sent_tokenize(text):
    return tokenizer.tokenize(text)


# nltk.download('punkt')
# from nltk.tokenize import sent_tokenize

bert_model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
DATA_FILE = "analysis_results.json"

def clean_text(text):
    return re.sub(r"[^a-zA-Z0-9\s?]", "", text).lower().strip()

def extract_text_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs]).strip()

def extract_text_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if not page_text:
                page_text = ocr_page(page)
            if page_text:
                text += page_text + "\n"
    return text.strip()

def remove_watermark_pdf(input_path):
    output_path = input_path.replace(".pdf", "_cleaned.pdf")
    reader = PdfReader(input_path)
    writer = PdfWriter()
    for page in reader.pages:
        if "/Annots" in page:
            page["/Annots"] = []
        writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_path

def remove_watermark_docx(input_path):
    doc = docx.Document(input_path)
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if "watermark" in paragraph.text.lower():
                paragraph.clear()
    output_path = input_path.replace(".docx", "_cleaned.docx")
    doc.save(output_path)
    return output_path

def ocr_page(page):
    image = page.to_image(resolution=300)
    img_pil = image.original.convert("RGB")
    text = pytesseract.image_to_string(img_pil)
    return text

def save_analysis_results(results):
    with open(DATA_FILE, "w") as file:
        json.dump(results, file)

def load_analysis_results():
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, "r") as file:
            return json.load(file)
    return []

def group_similar_questions_bert(questions, question_frequency, threshold=0.75):
    if not questions:
        return []
    unique_questions = list(set(questions))
    question_embeddings = bert_model.encode(unique_questions, convert_to_tensor=True)
    grouped_questions = []
    used_indices = set()
    question_map = defaultdict(list)
    for i, q1 in enumerate(unique_questions):
        if i in used_indices:
            continue
        group = [q1]
        used_indices.add(i)
        for j, q2 in enumerate(unique_questions):
            if i != j and j not in used_indices:
                similarity = util.pytorch_cos_sim(question_embeddings[i], question_embeddings[j]).item()
                if similarity > threshold:
                    group.append(q2)
                    used_indices.add(j)
        representative_question = group[0]
        question_map[representative_question] = group
    return [{
        "question": key,
        "similar_variants": value,
        "frequency": sum(question_frequency[q] for q in value)
    } for key, value in question_map.items()]

def analyze_questions(file_paths):
    all_questions = []
    question_frequency = defaultdict(int)
    for file_path in file_paths:
        if file_path.endswith(".pdf"):
            cleaned_path = remove_watermark_pdf(file_path)
            text = extract_text_pdf(cleaned_path)
        elif file_path.endswith(".docx"):
            cleaned_path = remove_watermark_docx(file_path)
            text = extract_text_docx(cleaned_path)
        else:
            continue
        sentences = sent_tokenize(text)
        cleaned_sentences = [clean_text(s.replace("\n", " ").strip()) for s in sentences]
        questions = [s for s in cleaned_sentences if s.endswith("?") and len(s.split()) > 3 and any(c.isalpha() for c in s)]
        for q in questions:
            question_frequency[q] += 1
        all_questions.extend(questions)
    grouped_questions = group_similar_questions_bert(all_questions, question_frequency)
    if not grouped_questions:
        return None
    save_analysis_results(grouped_questions)
    return grouped_questions

def get_ans_gpt(questions):
    from openai import OpenAI
    client = OpenAI(api_key="REMOVED_API_KEY")
    answers = []
    for question in questions:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a study assistant. Answer academic questions concisely."},
                {"role": "user", "content": question}
            ]
        )
        answer = response.choices[0].message.content
        answers.append({"question": question, "answer": answer})
    return answers
